"""
Word文档工具
"""
from agent.tools.registry import register_tool
from agent.tools.file_tools.path_utils import ensure_safe_path


def read_word(path: str) -> str:
    try:
        safe_path = ensure_safe_path(path)
    except ValueError as e:
        return f"错误: {str(e)}"
    try:
        from docx import Document
        doc = Document(safe_path)
        return "\n".join(p.text for p in doc.paragraphs)
    except ImportError:
        return "需要安装python-docx: pip install python-docx"
    except Exception as e:
        return f"错误: {str(e)}"


def write_word(path: str, content: str, mode: str = "create") -> str:
    try:
        safe_path = ensure_safe_path(path)
    except ValueError as e:
        return f"错误: {str(e)}"
    try:
        from docx import Document
        doc = Document() if mode == "create" else Document(safe_path)
        for line in content.split("\n"):
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            else:
                doc.add_paragraph(line)
        doc.save(safe_path)
        return f"Word文档已保存: {safe_path}"
    except ImportError:
        return "需要安装python-docx: pip install python-docx"
    except Exception as e:
        return f"错误: {str(e)}"


register_tool(name="read_word", description="读取Word(.docx)文件内容",
    parameters={"type": "object", "properties": {
        "path": {"type": "string", "description": "文件路径"}
    }, "required": ["path"]},
    handler=read_word, category="file")

register_tool(name="write_word", description="创建或修改Word(.docx)文件",
    parameters={"type": "object", "properties": {
        "path": {"type": "string", "description": "文件路径"},
        "content": {"type": "string", "description": "文件内容"},
        "mode": {"type": "string", "enum": ["create", "append"], "default": "create"}
    }, "required": ["path", "content"]},
    handler=write_word, category="file")
